from pathlib import Path

from docx import Document as DocxDocument

from jurisflow_parsers.types import ParsedDocument


def parse_docx(file_path: Path, mime_type: str) -> ParsedDocument:
    doc = DocxDocument(file_path)
    text = "\n".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
    return ParsedDocument(file_path=file_path, mime_type=mime_type, text=text)

